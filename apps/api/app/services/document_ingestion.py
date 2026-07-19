import hashlib
import io
import math
import re
import zipfile
from collections import Counter
from pathlib import Path, PurePosixPath
from uuid import uuid4

from docx import Document
from docx.document import Document as DocxDocument
from fastapi import UploadFile
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.core.security import detect_prompt_injection
from app.schemas.document import (
    DocumentExtractionResponse,
    DocumentSourceReference,
    DocumentType,
    ExtractedSection,
    SourceLocationType,
)

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
NORMALIZED_MEDIA_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
}
CONTROL_CHARACTERS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
MAX_DOCX_ENTRIES = 2_000
MAX_DOCX_EXPANDED_BYTES = 50 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 200


class DocumentIngestionError(Exception):
    def __init__(
        self,
        code: str,
        message: str,
        *,
        status_code: int = 400,
        details: dict[str, str | int] | None = None,
    ) -> None:
        super().__init__(message)
        self.code = code
        self.message = message
        self.status_code = status_code
        self.details = details or {}


async def read_upload_limited(upload: UploadFile, max_bytes: int) -> bytes:
    chunks: list[bytes] = []
    total = 0
    try:
        while chunk := await upload.read(64 * 1024):
            total += len(chunk)
            if total > max_bytes:
                raise DocumentIngestionError(
                    "FILE_TOO_LARGE",
                    "The uploaded file exceeds the configured size limit.",
                    status_code=413,
                    details={"max_bytes": max_bytes},
                )
            chunks.append(chunk)
    finally:
        await upload.close()
    return b"".join(chunks)


def ingest_document(
    data: bytes,
    filename: str,
    document_type: DocumentType,
    *,
    max_text_characters: int,
    max_pdf_pages: int,
) -> DocumentExtractionResponse:
    safe_filename = Path(filename or "upload").name
    suffix = Path(safe_filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise DocumentIngestionError(
            "UNSUPPORTED_FILE_TYPE",
            "Only PDF, DOCX, and TXT files are supported.",
            details={"extension": suffix or "none"},
        )
    if not data:
        raise DocumentIngestionError("EMPTY_FILE", "The uploaded file is empty.")

    document_id = str(uuid4())
    if suffix == ".txt":
        sections, references, warnings, confidence = _parse_txt(data, document_id)
    elif suffix == ".pdf":
        sections, references, warnings, confidence = _parse_pdf(data, document_id, max_pdf_pages)
    else:
        sections, references, warnings, confidence = _parse_docx(data, document_id)

    raw_text = "\n".join(reference.text for reference in references).strip()
    if not raw_text:
        raise DocumentIngestionError(
            "EXTRACTION_EMPTY",
            "No readable text could be extracted from the document.",
        )
    if len(raw_text) > max_text_characters:
        raise DocumentIngestionError(
            "EXTRACTED_TEXT_TOO_LARGE",
            "The extracted document text exceeds the configured analysis limit.",
            status_code=413,
            details={"max_characters": max_text_characters},
        )
    if len(raw_text) < 100:
        warnings.append("Very little text was extracted; review the document before analysis.")
    warnings.extend(detect_prompt_injection(raw_text, document_type.value.replace("_", " ")))

    return DocumentExtractionResponse(
        document_id=document_id,
        document_type=document_type,
        filename=safe_filename,
        media_type=NORMALIZED_MEDIA_TYPES[suffix],
        sha256=hashlib.sha256(data).hexdigest(),
        raw_text=raw_text,
        sections=sections,
        source_references=references,
        warnings=list(dict.fromkeys(warnings)),
        extraction_confidence=confidence,
        character_count=len(raw_text),
    )


def _clean_text(value: str) -> tuple[str, bool]:
    normalized = value.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = CONTROL_CHARACTERS.sub("", normalized)
    cleaned = "\n".join(line.rstrip() for line in cleaned.splitlines())
    return cleaned.strip(), cleaned != normalized.strip()


def _parse_txt(
    data: bytes, document_id: str
) -> tuple[list[ExtractedSection], list[DocumentSourceReference], list[str], float]:
    if data.startswith((b"%PDF-", b"PK\x03\x04")) or b"\x00" in data:
        raise DocumentIngestionError(
            "FILE_SIGNATURE_MISMATCH",
            "The file contents do not match a plain-text document.",
        )
    try:
        decoded = data.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise DocumentIngestionError(
            "INVALID_TEXT_ENCODING",
            "TXT files must use UTF-8 encoding.",
        ) from error

    text, removed_controls = _clean_text(decoded)
    warnings = ["Unsupported control characters were removed."] if removed_controls else []
    references = [
        DocumentSourceReference(
            id=f"{document_id}-line-{line_number}",
            text=line.strip(),
            location_type=SourceLocationType.LINE,
            line=line_number,
        )
        for line_number, line in enumerate(text.splitlines(), start=1)
        if line.strip()
    ]
    section = ExtractedSection(
        title="Document",
        text="\n".join(reference.text for reference in references),
        source_reference_ids=[reference.id for reference in references],
    )
    return [section], references, warnings, 0.99


def _parse_pdf(
    data: bytes, document_id: str, max_pages: int
) -> tuple[list[ExtractedSection], list[DocumentSourceReference], list[str], float]:
    if not data.startswith(b"%PDF-"):
        raise DocumentIngestionError(
            "FILE_SIGNATURE_MISMATCH",
            "The file contents do not match a PDF document.",
        )
    try:
        reader = PdfReader(io.BytesIO(data), strict=False)
        if reader.is_encrypted:
            raise DocumentIngestionError(
                "ENCRYPTED_DOCUMENT",
                "Password-protected PDF files are not supported.",
            )
        if len(reader.pages) > max_pages:
            raise DocumentIngestionError(
                "TOO_MANY_PAGES",
                "The PDF exceeds the configured page limit.",
                status_code=413,
                details={"max_pages": max_pages},
            )
        page_lines: list[list[str]] = []
        for page in reader.pages:
            page_text, _ = _clean_text(page.extract_text() or "")
            page_lines.append([line.strip() for line in page_text.splitlines() if line.strip()])
    except DocumentIngestionError:
        raise
    except (PdfReadError, ValueError, TypeError, KeyError) as error:
        raise DocumentIngestionError(
            "PDF_PARSE_FAILED",
            "The PDF could not be parsed safely.",
        ) from error

    warnings: list[str] = []
    repeated = _repeated_pdf_margin_lines(page_lines)
    if repeated:
        page_lines = [
            [line for line in lines if _margin_key(line) not in repeated] for lines in page_lines
        ]
        warnings.append("Repeated page headers or footers were removed.")

    references: list[DocumentSourceReference] = []
    sections: list[ExtractedSection] = []
    readable_pages = 0
    for page_number, lines in enumerate(page_lines, start=1):
        page_references: list[DocumentSourceReference] = []
        for line_number, line in enumerate(lines, start=1):
            reference = DocumentSourceReference(
                id=f"{document_id}-page-{page_number}-line-{line_number}",
                text=line,
                location_type=SourceLocationType.LINE,
                page=page_number,
                line=line_number,
            )
            references.append(reference)
            page_references.append(reference)
        if page_references:
            readable_pages += 1
            sections.append(
                ExtractedSection(
                    title=f"Page {page_number}",
                    text="\n".join(reference.text for reference in page_references),
                    source_reference_ids=[reference.id for reference in page_references],
                )
            )

    page_count = len(page_lines)
    if readable_pages < page_count:
        warnings.append(
            f"No readable text was found on {page_count - readable_pages} PDF page(s); "
            "OCR is not performed."
        )
    confidence = min(0.96, 0.55 + (0.4 * readable_pages / max(page_count, 1)))
    return sections, references, warnings, round(confidence, 2)


def _margin_key(value: str) -> str:
    return re.sub(r"\d+", "#", re.sub(r"\s+", " ", value.lower())).strip()


def _repeated_pdf_margin_lines(page_lines: list[list[str]]) -> set[str]:
    if len(page_lines) < 3:
        return set()
    candidates = [
        _margin_key(line)
        for lines in page_lines
        for line in ([lines[0]] if lines else []) + ([lines[-1]] if len(lines) > 1 else [])
        if len(line) <= 160
    ]
    threshold = math.ceil(len(page_lines) * 0.6)
    return {line for line, count in Counter(candidates).items() if count >= threshold}


def _validate_docx_archive(data: bytes) -> list[str]:
    if not data.startswith(b"PK\x03\x04") or not zipfile.is_zipfile(io.BytesIO(data)):
        raise DocumentIngestionError(
            "FILE_SIGNATURE_MISMATCH",
            "The file contents do not match a DOCX document.",
        )
    warnings: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ENTRIES:
                raise DocumentIngestionError(
                    "UNSAFE_DOCX_ARCHIVE", "The DOCX archive contains too many entries."
                )
            names = {entry.filename for entry in entries}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise DocumentIngestionError(
                    "INVALID_DOCX_STRUCTURE", "The DOCX package is missing required content."
                )
            for entry in entries:
                path = PurePosixPath(entry.filename)
                if path.is_absolute() or ".." in path.parts:
                    raise DocumentIngestionError(
                        "UNSAFE_DOCX_ARCHIVE", "The DOCX archive contains an unsafe path."
                    )
            expanded_size = sum(entry.file_size for entry in entries)
            compressed_size = sum(max(entry.compress_size, 1) for entry in entries)
            if (
                expanded_size > MAX_DOCX_EXPANDED_BYTES
                or expanded_size / compressed_size > MAX_DOCX_COMPRESSION_RATIO
            ):
                raise DocumentIngestionError(
                    "UNSAFE_DOCX_ARCHIVE",
                    "The DOCX archive expands beyond the safe processing limit.",
                )
            if any(name.lower().endswith("vbaproject.bin") for name in names):
                raise DocumentIngestionError(
                    "MACRO_ENABLED_DOCUMENT",
                    "Macro-enabled Office documents are not supported.",
                )
            relationships = (
                archive.read("word/_rels/document.xml.rels")
                if "word/_rels/document.xml.rels" in names
                else b""
            )
            if b'TargetMode="External"' in relationships:
                warnings.append(
                    "External document relationships were detected and were not accessed."
                )
    except DocumentIngestionError:
        raise
    except (zipfile.BadZipFile, OSError, KeyError) as error:
        raise DocumentIngestionError(
            "DOCX_PARSE_FAILED", "The DOCX package could not be parsed safely."
        ) from error
    return warnings


def _parse_docx(
    data: bytes, document_id: str
) -> tuple[list[ExtractedSection], list[DocumentSourceReference], list[str], float]:
    warnings = _validate_docx_archive(data)
    try:
        document: DocxDocument = Document(io.BytesIO(data))
    except (ValueError, KeyError, OSError, zipfile.BadZipFile) as error:
        raise DocumentIngestionError(
            "DOCX_PARSE_FAILED", "The DOCX content could not be parsed safely."
        ) from error

    references: list[DocumentSourceReference] = []
    section_groups: list[tuple[str, list[DocumentSourceReference]]] = []
    current_title = "Document"
    current_references: list[DocumentSourceReference] = []

    def flush_section() -> None:
        nonlocal current_references
        if current_references:
            section_groups.append((current_title, current_references))
            current_references = []

    for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
        text, _ = _clean_text(paragraph.text)
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        if style_name.startswith("heading"):
            flush_section()
            current_title = text[:120]
        reference = DocumentSourceReference(
            id=f"{document_id}-paragraph-{paragraph_number}",
            text=text,
            location_type=SourceLocationType.PARAGRAPH,
            paragraph=paragraph_number,
        )
        references.append(reference)
        current_references.append(reference)

    for table_number, table in enumerate(document.tables, start=1):
        for row_number, row in enumerate(table.rows, start=1):
            text, _ = _clean_text(" | ".join(cell.text.strip() for cell in row.cells))
            if not text:
                continue
            reference = DocumentSourceReference(
                id=f"{document_id}-table-{table_number}-row-{row_number}",
                text=text,
                location_type=SourceLocationType.TABLE_ROW,
                table=table_number,
                row=row_number,
            )
            references.append(reference)
            current_references.append(reference)
    flush_section()

    sections = [
        ExtractedSection(
            title=title,
            text="\n".join(reference.text for reference in group),
            source_reference_ids=[reference.id for reference in group],
        )
        for title, group in section_groups
    ]
    return sections, references, warnings, 0.97
